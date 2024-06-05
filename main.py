from docx import Document
from docx.shared import Pt
import json
import requests
from docx.shared import Inches
from io import BytesIO
import urllib.request
from datetime import datetime


doc = Document('##')

current_datetime = datetime.now()

current_date = datetime.now().strftime("%d-%m-%Y")

placeholder = 'Date: 02-06-2023'


for paragraph in doc.paragraphs:
    if placeholder in paragraph.text:
        updated_text = paragraph.text.replace(placeholder, f"Date: {current_date}")
        paragraph.text = updated_text
        paragraph.runs[0].font.size = Pt(12) 




order={0: 'eventCode', 1: 'eventName', 2: 'eventDate', 3: 'eventTime', 4: 'eventAbout', 5: 'eventFee', 6: 'eventFormLink', 7: 'eventLocation', 9: 'organizerName', 10: 'eventAddInfo'}



data= {
  "coverPhoto": "https://i.imgur.com/1lH20gT.jpg",
  "eventAbout": "Dance Fest 2023 is a celebration of dance in all its forms. From contemporary to hip-hop, ballet to breakdance, our talented performers will take the stage and captivate you with their stunning choreography and breathtaking moves. Prepare to be enthralled as they push the boundaries of creativity and showcase the beauty and power of dance.\n\nWhether you're a passionate dancer yourself or simply appreciate the artistry, Groove Fest 2023 promises to be an unforgettable experience. Immerse yourself in the vibrant atmosphere as the performers transport you to a world of rhythm and expression.\n\nIn addition to the spectacular performances, Groove Fest 2023 offers interactive workshops led by renowned dance instructors. Join us to learn new techniques, enhance your skills, and explore various dance styles. It's an opportunity to connect with fellow dance enthusiasts, share your love for dance, and ignite your passion for movement.\n\nSo mark your calendars and spread the word! Groove Fest 2023 is not to be missed. Come and be a part of this incredible celebration of dance, where artistry and creativity collide in an explosion of talent and inspiration.\n\nStay tuned for further updates and ticket information. Get ready to groove like never before at Dance Fest 2023!",
  "eventAddInfo": "• Dance like a champ\n• No rules\n• Winner gets ₹5000\n• Call 8919798735",
  "eventCode": "DANC_DAST_14062023_1015",
  "eventDate": "14-06-2023",
  "eventFee": 0,
  "eventFormLink": "",
  "eventLocation": "Sangeeth Auditorium",
  "eventName": "Dance Fest 2023",
  "eventOrganizers": {
    "organizerBio": "Dancers at IARE",
    "organizerLink": "https://gdsciare.club",
    "organizerLogo": "https://i.imgur.com/ezUdM6U.jpg",
    "organizerName": "Dancing Club"
  },
  "eventPoster": {
    "posterLink": "https://i.imgur.com/4k4O9U3.jpg"
  },
  "eventTag1": "Dance",
  "eventTag2": "Fun",
  "eventTime": "10:15"
}


def mapping(row_index,order,data):
    column_index=1
    table = doc.tables[0] 
    cell = table.cell(row_index, column_index)
    paragraph = cell.paragraphs[0]
    paragraph_format = paragraph.paragraph_format
    run = paragraph.runs[0]
    font = run.font
    spacing_before = Pt(0)
    spacing_after = Pt(0)

    if row_index==8:
          pass
    elif row_index==9:
          new_text=str(data['eventOrganizers']['organizerName'])
          cell.text = new_text
    else:
          new_text = str(data[order[row_index]])
          cell.text = new_text

    new_paragraph = cell.paragraphs[0]
    new_paragraph.alignment = paragraph_format.alignment
    new_run = new_paragraph.runs[0]
    new_run.font.size = font.size
    new_run.font.name = font.name
    

    for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = spacing_before
            paragraph.paragraph_format.space_after = spacing_after

for i in range(10):
     mapping(i,order,data)

image_url = '##'

response = requests.get(image_url)
image_data = BytesIO(response.content)
doc.add_picture(image_data,width=Inches(12))

doc.save('##')

